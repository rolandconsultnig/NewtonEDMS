"""Tests for Microsoft Office Integration: WOPI, Desktop URI, Add-in Manifests, Properties & Templates."""
import io
import json
import base64
import pytest
from pathlib import Path
from fastapi.testclient import TestClient

from app.main import app
from app.wopi import generate_wopi_token
from app.office_meta import (
    get_office_properties,
    update_office_properties,
    merge_word_template,
    merge_excel_template,
    export_documents_to_excel,
    export_documents_to_word,
)


@pytest.fixture
def auth_client(client, admin_token):
    client.headers.update({"Authorization": f"Bearer {admin_token}"})
    return client


@pytest.fixture
def sample_docx(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_heading("Contract Agreement", level=0)
    doc.add_paragraph("This agreement is between {{client}} and {{provider}} on date {{date}}.")
    file_path = tmp_path / "sample.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def sample_xlsx(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Client Name: {{client}}"
    ws["A2"] = "Total: {{amount}}"
    file_path = tmp_path / "sample.xlsx"
    wb.save(str(file_path))
    wb.close()
    return file_path


def test_office_meta_properties(sample_docx):
    # Test reading & updating OpenXML properties
    props = get_office_properties(sample_docx)
    assert isinstance(props, dict)
    
    updated = update_office_properties(
        sample_docx,
        {
            "title": "Master Services Agreement",
            "author": "Chief Legal Officer",
            "subject": "Commercial Contract",
            "keywords": "contract, legal, 2026",
            "comments": "Reviewed and approved",
        },
    )
    assert updated is True
    
    reloaded = get_office_properties(sample_docx)
    assert reloaded["title"] == "Master Services Agreement"
    assert reloaded["author"] == "Chief Legal Officer"
    assert reloaded["keywords"] == "contract, legal, 2026"


def test_word_and_excel_template_merging(sample_docx, sample_xlsx, tmp_path):
    # Test Word template merge
    docx_out = tmp_path / "merged.docx"
    merge_word_template(
        sample_docx,
        docx_out,
        {"client": "Globex Corp", "provider": "Newton Industries", "date": "2026-08-20"},
    )
    assert docx_out.exists()
    
    import docx
    res_doc = docx.Document(str(docx_out))
    full_text = " ".join(p.text for p in res_doc.paragraphs)
    assert "Globex Corp" in full_text
    assert "Newton Industries" in full_text
    assert "{{client}}" not in full_text
    
    # Test Excel template merge
    xlsx_out = tmp_path / "merged.xlsx"
    merge_excel_template(sample_xlsx, xlsx_out, {"client": "Acme Co", "amount": "$10,000"})
    assert xlsx_out.exists()
    
    import openpyxl
    wb = openpyxl.load_workbook(str(xlsx_out))
    ws = wb.active
    assert ws["A1"].value == "Client Name: Acme Co"
    assert ws["A2"].value == "Total: $10,000"
    wb.close()


def test_export_documents_to_excel_and_word():
    docs_data = [
        {"id": 1, "name": "Annual Report.docx", "folder_name": "Finance", "version": "1.0", "status": "approved", "size": 10240, "tags": ["annual", "finance"]},
        {"id": 2, "name": "Budget 2026.xlsx", "folder_name": "Finance", "version": "2.1", "status": "draft", "size": 20480, "tags": ["budget"]},
    ]
    excel_bytes = export_documents_to_excel(docs_data, title="Financial Documents")
    assert len(excel_bytes) > 100
    assert excel_bytes[:2] == b"PK"  # Zip/OpenXML magic bytes
    
    word_bytes = export_documents_to_word(docs_data, title="Executive Dossier")
    assert len(word_bytes) > 100
    assert word_bytes[:2] == b"PK"


def test_office_addin_manifests(auth_client):
    # XML Manifest
    res_xml = auth_client.get("/api/office/addin/manifest.xml")
    assert res_xml.status_code == 200
    assert "OfficeApp" in res_xml.text
    assert "NewtonEDMS" in res_xml.text
    assert "taskpane.html" in res_xml.text
    
    # JSON Manifest
    res_json = auth_client.get("/api/office/addin/manifest.json")
    assert res_json.status_code == 200
    data = res_json.json()
    assert data["name"]["short"] == "NewtonEDMS"
    assert data["extensions"][0]["runtimes"][0]["id"] == "NewtonTaskpaneRuntime"
    
    # Add-in Info
    res_info = auth_client.get("/api/office/addin/info")
    assert res_info.status_code == 200
    assert "Word" in res_info.json()["supported_apps"]


def test_wopi_protocol_and_desktop_integration(auth_client, sample_docx, root_folder_id):
    # Upload a document first
    with open(sample_docx, "rb") as f:
        res_upload = auth_client.post(
            "/api/documents",
            data={"folder_id": root_folder_id},
            files={"file": ("contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert res_upload.status_code == 200, res_upload.text
    doc_id = res_upload.json()["id"]
    
    # 1. Desktop Launch endpoint
    res_launch = auth_client.get(f"/api/office/desktop-launch/{doc_id}")
    assert res_launch.status_code == 200
    launch_data = res_launch.json()
    assert "ms-word:ofe|u|" in launch_data["protocol_uri"]
    assert launch_data["app_name"] == "Microsoft Word"
    
    # 2. WOPI Session endpoint
    res_session = auth_client.get(f"/api/office/wopi/session/{doc_id}")
    assert res_session.status_code == 200
    session_data = res_session.json()
    token = session_data["access_token"]
    assert token
    
    # 3. WOPI CheckFileInfo
    res_check = auth_client.get(f"/wopi/files/{doc_id}?access_token={token}")
    assert res_check.status_code == 200
    info = res_check.json()
    assert info["BaseFileName"] == "contract.docx"
    assert info["SupportsLocks"] is True
    assert info["UserCanWrite"] is True
    assert info["BreadcrumbBrandName"] == "NewtonEDMS"
    
    # 4. WOPI GetFile
    res_get = auth_client.get(f"/wopi/files/{doc_id}/contents?access_token={token}")
    assert res_get.status_code == 200
    assert len(res_get.content) > 0
    
    # 5. WOPI Lock operation
    lock_id = "wopi-lock-session-abc"
    res_lock = auth_client.post(
        f"/wopi/files/{doc_id}?access_token={token}",
        headers={"X-WOPI-Override": "LOCK", "X-WOPI-Lock": lock_id},
    )
    assert res_lock.status_code == 200
    
    # 6. WOPI Get Lock
    res_get_lock = auth_client.post(
        f"/wopi/files/{doc_id}?access_token={token}",
        headers={"X-WOPI-Override": "GET_LOCK"},
    )
    assert res_get_lock.status_code == 200
    assert res_get_lock.headers["X-WOPI-Lock"] == lock_id
    
    # 7. WOPI PutFile
    new_bytes = b"Updated Word document content via Office Online"
    res_put = auth_client.post(
        f"/wopi/files/{doc_id}/contents?access_token={token}",
        data=new_bytes,
        headers={"X-WOPI-Lock": lock_id},
    )
    assert res_put.status_code == 200
    assert "ItemVersion" in res_put.json()
    
    # 8. WOPI Unlock
    res_unlock = auth_client.post(
        f"/wopi/files/{doc_id}?access_token={token}",
        headers={"X-WOPI-Override": "UNLOCK", "X-WOPI-Lock": lock_id},
    )
    assert res_unlock.status_code == 200
    
    # 9. Properties endpoint
    res_props = auth_client.get(f"/api/office/properties/{doc_id}")
    assert res_props.status_code == 200


def test_office_template_merge_endpoint(auth_client, sample_docx, root_folder_id):
    # Upload template
    with open(sample_docx, "rb") as f:
        res_upload = auth_client.post(
            "/api/documents",
            data={"folder_id": root_folder_id},
            files={"file": ("TemplateContract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert res_upload.status_code == 200, res_upload.text
    template_id = res_upload.json()["id"]
    
    # Merge template
    res_merge = auth_client.post(
        f"/api/office/templates/{template_id}/merge",
        json={
            "target_name": "Final_Acme_Contract.docx",
            "context": {"client": "Acme Inc.", "provider": "NewtonEDMS Ltd."},
        },
    )
    assert res_merge.status_code == 200
    data = res_merge.json()
    assert data["status"] == "success"
    assert data["name"] == "Final_Acme_Contract.docx"
    assert data["new_document_id"] > 0


def test_outlook_archive_endpoint(auth_client):
    res = auth_client.post(
        "/api/office/outlook/archive",
        json={
            "subject": "Q3 Enterprise Partnership Agreement",
            "from_address": "partner@acme.com",
            "from_name": "John Partner",
            "sent_date": "2026-08-20T16:00:00Z",
            "body_html": "<p>Please find attached the signed contract terms.</p>",
            "tags": ["outlook", "partner", "signed"],
            "attachments": [
                {
                    "filename": "Terms.txt",
                    "content_base64": base64.b64encode(b"Confidential Agreement Terms").decode("ascii"),
                }
            ],
        },
    )
    assert res.status_code == 200
    data = res.json()
    assert data["status"] == "success"
    assert data["email_document_id"] > 0
    assert data["attachments_count"] == 1
